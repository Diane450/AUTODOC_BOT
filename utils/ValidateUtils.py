from aiogram.types import Document
from urllib.parse import urlparse
import zipfile
import re
from docxtpl import DocxTemplate
from jinja2.exceptions import TemplateSyntaxError
from docx import Document
import html

class ValidateUtils:

    def looks_like_gsheets(self, url: str) -> bool:
        if "docs.google.com/spreadsheets" in url:
            return True
        parsed = urlparse(url)
        if parsed.scheme in ("http", "https") and parsed.path.lower().endswith(".xlsx"):
            return True
        return False
    

    def is_docx(self, doc:Document):
        return doc.file_name.lower().endswith(".docx")

    def check_template_syntax(self, file_path: str) -> tuple[bool, str]:
        try:
            doc = DocxTemplate(file_path)
            doc.get_undeclared_template_variables()
            return True, ""
            
        except TemplateSyntaxError as e:
            user_friendly_reason = "Синтаксическая ошибка в плейсхолдере."
            if "expected token" in e.message or "expected" in e.message:
                user_friendly_reason = "Опечатка внутри скобок (возможно, лишний символ или не хватает фигурной скобки )."
            elif "unexpected end of template" in e.message:
                user_friendly_reason = "Вы забыли закрыть тег двойными скобками '}}'."
            
            error_context = ""
            try:
                word_doc = Document(file_path)
                
                texts = [p.text for p in word_doc.paragraphs]
                for table in word_doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            texts.extend([p.text for p in cell.paragraphs])
                
                suspicious_paragraphs =[]
                
                for text in texts:
                    if '{' in text or '}' in text:
                        clean_text = re.sub(r'\{\{.*?\}\}', '', text)
                        clean_text = re.sub(r'\{%.*?%\}', '', clean_text)
                        clean_text = re.sub(r'\{#.*?#\}', '', clean_text)
                        
                        if '{' in clean_text or '}' in clean_text:
                            if len(text.strip()) > 0:
                                suspicious_paragraphs.append(text.strip())
                
                if suspicious_paragraphs:
                    raw_error_context = "\n".join(suspicious_paragraphs[:3])
                    error_context = html.escape(raw_error_context)
                    
            except Exception as inner_e:
                print(f"ОШИБКА ИЗВЛЕЧЕНИЯ ТЕКСТА: {inner_e}")
                
            if not error_context:
                error_context = "⚠️ Не удалось точно извлечь фрагмент текста. Проверьте все поля {{ ... }} глазами."

            error_msg = (
                f"❌ <b>Обнаружена ошибка в вашем шаблоне!</b>\n\n"
                f"<b>Проблема:</b> {user_friendly_reason}\n\n"
                f"<b>Где искать ошибку (подозрительный текст):</b>\n"
                f"<code>{error_context}</code>\n\n"
            )
            return False, error_msg
            
        except Exception as e:
            return False, f"❌ Не удалось прочитать шаблон. Возможно, файл поврежден.\nДетали: {str(e)}"